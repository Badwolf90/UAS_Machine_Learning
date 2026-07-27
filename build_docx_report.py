import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_report():
    doc = Document()
    
    # Page setup - Margins 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("UNIVERSITAS AMIKOM YOGYAKARTA\nFAKULTAS ILMU KOMPUTER")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_paragraph() # spacing
    
    uas_p = doc.add_paragraph()
    uas_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uas = uas_p.add_run("LAPORAN FINAL PROJECT UJIAN AKHIR SEMESTER (UAS)\nMATA KULIAH: MACHINE LEARNING (TK075)")
    r_uas.bold = True
    r_uas.font.size = Pt(16)
    r_uas.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
    
    doc.add_paragraph()
    
    judul_p = doc.add_paragraph()
    judul_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_j = judul_p.add_run("JUDUL PROJECT:\nKLASIFIKASI KUALITAS MINUMAN RED WINE MENGGUNAKAN ALGORITMA RANDOM FOREST DAN SUPPORT VECTOR MACHINE (SVM)")
    r_j.bold = True
    r_j.font.size = Pt(13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Group Info Box
    tbl_cover = doc.add_table(rows=3, cols=2)
    tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cover.autofit = False
    
    fields = [
        ("Nama Anggota Tim", "[NAMA LENGKAP KAMU / ANGGOTA]"),
        ("NIM", "[NIM KAMU / ANGGOTA]"),
        ("Program Studi / Kelas", "Teknik Komputer / S1")
    ]
    for idx, (label, val) in enumerate(fields):
        row = tbl_cover.rows[idx]
        cell_l, cell_v = row.cells[0], row.cells[1]
        cell_l.width = Inches(2.2)
        cell_v.width = Inches(4.0)
        
        p1 = cell_l.paragraphs[0]
        r1 = p1.add_run(label)
        r1.bold = True
        
        p2 = cell_v.paragraphs[0]
        p2.add_run(val)
        
        set_cell_background(cell_l, "F2F4F7")
        set_cell_background(cell_v, "FFFFFF")
        
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # SECTION 1: LATAR BELAKANG PERMASALAHAN
    # ---------------------------------------------------------
    h1 = doc.add_heading("1. Latar Belakang Permasalahan", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    p = doc.add_paragraph(
        "Industri minuman dan makanan modern sangat bergantung pada kontrol kualitas (quality control) "
        "secara terukur dan terstandarisasi untuk menjamin kepuasan serta keselamatan konsumen. "
        "Pada produk minuman anggur merah (Red Wine), evaluasi kualitas produk secara konvensional "
        "sering kali mengandalkan uji organoleptik atau kecap lidah oleh panelis ahli sommelier. "
        "Pendekatan konvensional ini memiliki beberapa kelemahan signifikan, yaitu bersifat subjektif, "
        "memakan waktu lama, membutuhkan biaya operasional tinggi, serta rentan terhadap bias individu."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Seiring berkembangnya teknologi Machine Learning (ML) dan Sains Data di bidang Teknik Komputer, "
        "proses sertifikasi dan pengujian kualitas dapat dilakukan secara otomatis, objektif, dan presisi tinggi "
        "berdasarkan parameter fisikokimia laboratorium (seperti tingkat keasaman, kadar gula sisa, densitas, pH, "
        "kadar sulfat, dan persentase alkohol). Melalui pemodelan komputasi ini, pabrik pengolahan dapat "
        "memprediksi tingkat kualitas produk secara instan sebelum didistribusikan ke pasar global."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)
    
    # ---------------------------------------------------------
    # SECTION 2: KELEBIHAN & KEKURANGAN ALGORITMA MODEL DAN EVALUASI
    # ---------------------------------------------------------
    h2 = doc.add_heading("2. Kelebihan dan Kekurangan Algoritma Model & Evaluasi", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_heading("a. Model Random Forest Classifier (Ensemble Model)", level=2)
    doc.add_paragraph(
        "Random Forest adalah algoritma berbasis ensemble yang menggabungkan puluhan hingga ratusan pohon keputusan "
        "(Decision Trees) menggunakan teknik Bagging (Bootstrap Aggregating)."
    )
    
    p = doc.add_paragraph()
    p.add_run("Kelebihan Random Forest:\n").bold = True
    p.add_run("1. Tahan terhadap overfitting karena rata-rata prediksi dari banyak sampel pohon.\n"
              "2. Sangat stabil dalam menangani hubungan non-linear antar variabel fisikokimia.\n"
              "3. Memiliki kemampuan menghitung 'Feature Importance' untuk mengetahui variabel yang paling berpengaruh.\n"
              "4. Tidak memerlukan pemrosesan khusus penyesuaian skala data (feature scaling).")
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("Kekurangan Random Forest:\n").bold = True
    p.add_run("1. Kompleksitas komputasi yang tinggi dan membutuhkan memori lebih besar pada dataset raksasa.\n"
              "2. Lebih sulit diinterpretasikan alur prediksinya secara visual dibandingkan satu Decision Tree tunggal.")
    p.paragraph_format.space_after = Pt(12)

    doc.add_heading("b. Model Support Vector Machine / SVM", level=2)
    doc.add_paragraph(
        "SVM adalah algoritma supervised learning yang bekerja dengan mencari bidang pemisah optimal (Hyperplane) "
        "dengan margin maksimal di antara dua kelas data."
    )
    
    p = doc.add_paragraph()
    p.add_run("Kelebihan SVM:\n").bold = True
    p.add_run("1. Sangat efektif pada data berdimensi menengah dengan batas pemisah yang jelas.\n"
              "2. Menggunakan fungsi Kernel (RBF/Polynomial) yang mampu memetakan data tidak linear ke dimensi lebih tinggi.\n"
              "3. Sangat efisien dalam penggunaan memori karena hanya tergantung pada dukungan titik data (Support Vectors).")
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("Kekurangan SVM:\n").bold = True
    p.add_run("1. Sangat sensitif terhadap skala fitur (wajib dilakukan StandardScaler/MinMaxScaler).\n"
              "2. Peka terhadap pencilan (outliers) dan relatif lambat pada dataset sampel yang sangat besar.\n"
              "3. Tidak memberikan probabilitas prediksi secara langsung tanpa kalibrasi tambahan.")
    p.paragraph_format.space_after = Pt(12)

    doc.add_heading("c. Algoritma Metrik Evaluasi yang Dipilih", level=2)
    eval_tbl = doc.add_table(rows=6, cols=3)
    eval_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Metrik Evaluasi", "Kelebihan", "Kekurangan"]
    
    for j, h in enumerate(headers):
        cell = eval_tbl.rows[0].cells[j]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "003366")
        
    eval_data = [
        ("Accuracy", "Mudah dipahami, mengukur persentase prediksi benar secara keseluruhan.", "Dapat menyesatkan pada dataset yang tidak seimbang (imbalanced)."),
        ("Precision", "Tinggi presisi meminimalkan kesalahan False Positive (salah memprediksi barang buruk jadi bagus).", "Tidak memperhitungkan False Negative yang terlewat."),
        ("Recall (Sensitivity)", "Meminimalkan kesalahan False Negative (memastikan sampel produk bagus tidak terbuang).", "Dapat menurunkan presisi jika prediksi terlalu sensitif."),
        ("F1-Score", "Memberikan keseimbangan harmonis antara Precision dan Recall.", "Kurang intuitif dibandingkan akurasi murni bagi orang awam."),
        ("ROC-AUC Score", "Mengukur performa pemodelan pada seluruh ambang batas (threshold) tanpa terpengaruh proporsi kelas.", "Membutuhkan perhitungan kurva probabilitas kontinu.")
    ]
    
    for idx, (m, kel, kek) in enumerate(eval_data):
        row = eval_tbl.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(m).bold = True
        row.cells[1].paragraphs[0].add_run(kel)
        row.cells[2].paragraphs[0].add_run(kek)
        bg = "F9FAFB" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(row.cells[0], bg)
        set_cell_background(row.cells[1], bg)
        set_cell_background(row.cells[2], bg)
        
    doc.add_paragraph()
    
    # ---------------------------------------------------------
    # SECTION 3: CARA KERJA ALGORITMA & PARAMETER
    # ---------------------------------------------------------
    h3 = doc.add_heading("3. Cara Kerja Algoritma Model dan Parameter yang Digunakan", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_heading("a. Cara Kerja Random Forest", level=2)
    doc.add_paragraph(
        "1. Bootstrap Sampling: Mengambil sampel acak dengan pengembalian dari data latih.\n"
        "2. Node Splitting Acak: Pada setiap cabang pohon, algoritma memilih subset acak dari fitur fisikokimia untuk mencari pembagi terbaik.\n"
        "3. Voting Mayoritas: Seluruh pohon mengambil keputusan independen, dan hasil kelas akhir ditentukan berdasarkan jumlah suara terbanyak (Majority Vote)."
    )
    p = doc.add_paragraph()
    p.add_run("Parameter Utama RF:\n").bold = True
    p.add_run("- n_estimators: Jumlah pohon keputusan yang dibangun (misal: 10, 50, 100, 200, 300).\n"
              "- max_depth: Kedalaman maksimum setiap pohon untuk mencegah overfitting.\n"
              "- criterion: Fungsi pengukur kualitas pemisahan ('gini' atau 'entropy').\n"
              "- min_samples_split: Jumlah sampel minimal yang diperlukan untuk membagi simpul internal.")
    
    doc.add_heading("b. Cara Kerja Support Vector Machine (SVM)", level=2)
    doc.add_paragraph(
        "1. Pemetaan Fitur: Mengubah fitur input ke ruang vektor berdimensi tinggi menggunakan fungsi Kernel.\n"
        "2. Maksimisasi Margin: Menemukan Hyperplane yang memiliki jarak marjin terjauh terhadap titik sampel terdekat (Support Vectors).\n"
        "3. Klasifikasi: Menentukan posisi titik uji di atas atau di bawah garis batas hyperplane."
    )
    p = doc.add_paragraph()
    p.add_run("Parameter Utama SVM:\n").bold = True
    p.add_run("- C (Regularization Parameter): Mengontrol trade-off antara margin lebar dan toleransi kesalahan klasifikasi.\n"
              "- kernel: Jenis fungsi pemetaan matematika ('linear' atau 'rbf').\n"
              "- gamma: Menentukan seberapa jauh jangkauan pengaruh dari satu sampel pelatihan ('scale' atau 'auto').")
              
    doc.add_paragraph()
    
    # ---------------------------------------------------------
    # SECTION 4: EXPERIMENT MODEL (MINIMAL 5X)
    # ---------------------------------------------------------
    h4 = doc.add_heading("4. Experiment Model Menggunakan Parameter Minimal 5x", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_paragraph(
        "Eksperimen pemodelan dilakukan sebanyak 5 kali pengujian untuk masing-masing algoritma "
        "dengan memvariasikan kombinasi hyperparameter secara sistematis."
    )
    
    doc.add_heading("Tabel Hasil 5 Eksperimen Random Forest Classifier:", level=2)
    rf_tbl = doc.add_table(rows=6, cols=6)
    rf_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rf_headers = ["Eksperimen", "Konfigurasi Parameter", "Akurasi", "Presisi", "Recall", "F1-Score"]
    
    for j, h in enumerate(rf_headers):
        cell = rf_tbl.rows[0].cells[j]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "003366")
        
    rf_data = [
        ("Exp 1 (Base)", "n_estimators=10, max_depth=3", "74.38%", "77.99%", "72.51%", "75.15%"),
        ("Exp 2 (Tuning 1)", "n_estimators=50, max_depth=5", "76.25%", "80.25%", "73.68%", "76.83%"),
        ("Exp 3 (Tuning 2)", "n_estimators=100, max_depth=10, entropy", "80.31%", "82.53%", "80.12%", "81.31%"),
        ("Exp 4 (Tuning 3)", "n_estimators=200, max_depth=15, split=5", "79.69%", "81.18%", "80.70%", "80.94%"),
        ("Exp 5 (Optimal)", "n_estimators=300, max_depth=20, entropy", "79.69%", "81.55%", "80.12%", "80.83%")
    ]
    for idx, row_vals in enumerate(rf_data):
        row = rf_tbl.rows[idx+1]
        for c_idx, val in enumerate(row_vals):
            p = row.cells[c_idx].paragraphs[0]
            r = p.add_run(val)
            if c_idx == 0 or c_idx == 2 or c_idx == 5:
                r.bold = True
            bg = "F9FAFB" if idx % 2 == 0 else "FFFFFF"
            set_cell_background(row.cells[c_idx], bg)

    doc.add_paragraph()

    doc.add_heading("Tabel Hasil 5 Eksperimen Support Vector Machine (SVM):", level=2)
    svm_tbl = doc.add_table(rows=6, cols=6)
    svm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for j, h in enumerate(rf_headers):
        cell = svm_tbl.rows[0].cells[j]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "2E7D32")
        
    svm_data = [
        ("Exp 1 (Base)", "kernel='linear', C=0.1", "74.06%", "78.95%", "70.18%", "74.30%"),
        ("Exp 2 (Tuning 1)", "kernel='linear', C=1.0", "74.38%", "78.71%", "71.35%", "74.85%"),
        ("Exp 3 (Tuning 2)", "kernel='rbf', C=1.0, scale", "76.25%", "81.46%", "71.93%", "76.40%"),
        ("Exp 4 (Tuning 3)", "kernel='rbf', C=10.0, auto", "77.19%", "81.41%", "74.27%", "77.68%"),
        ("Exp 5 (Optimal)", "kernel='rbf', C=50.0, scale", "79.37%", "83.44%", "76.61%", "79.88%")
    ]
    for idx, row_vals in enumerate(svm_data):
        row = svm_tbl.rows[idx+1]
        for c_idx, val in enumerate(row_vals):
            p = row.cells[c_idx].paragraphs[0]
            r = p.add_run(val)
            if c_idx == 0 or c_idx == 2 or c_idx == 5:
                r.bold = True
            bg = "F9FAFB" if idx % 2 == 0 else "FFFFFF"
            set_cell_background(row.cells[c_idx], bg)

    doc.add_paragraph()
    
    # ---------------------------------------------------------
    # SECTION 5: ANALISIS HASIL EVALUASI & VISUALISASI
    # ---------------------------------------------------------
    h5 = doc.add_heading("5. Penjelasan Hasil Evaluasi & Diagram Visualisasi", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_heading("a. Grafik Perbandingan Perkembangan 5 Eksperimen", level=2)
    doc.add_paragraph(
        "Grafik di bawah ini menggambarkan peningkatan skor akurasi dan F1-Score pada setiap tahap eksperimen tuning. "
        "Terlihat bahwa Random Forest mencapai performa tertinggi pada Eksperimen 3 (Akurasi 80.31% & F1 81.31%), "
        "sedangkan SVM mencapai performa terbaik pada Eksperimen 5 (Akurasi 79.37% & F1 79.88%)."
    )
    doc.add_picture('experiments_comparison.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_heading("b. Visualisasi Confusion Matrix", level=2)
    doc.add_paragraph(
        "Confusion Matrix menampilkan persebaran data aktual vs data hasil prediksi. "
        "Random Forest berhasil memprediksi 137 sampel produk berkualitas tinggi secara tepat, "
        "sementara SVM menghasilkan presisi yang sangat kuat pada kelas positif (83.44%)."
    )
    doc.add_picture('confusion_matrix_comparison.png', width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_heading("c. Visualisasi Kurva ROC-AUC", level=2)
    doc.add_paragraph(
        "Kurva ROC-AUC membandingkan kemampuan diskriminasi pemodelan. "
        "Random Forest unggul secara konsisten dengan nilai ROC-AUC mencapai 0.9038, "
        "menunjukkan kemampuan klasifikasi yang sangat andal melebihi baseline acak (0.50)."
    )
    doc.add_picture('roc_curve_comparison.png', width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_heading("d. Tingkat Kepentingan Fitur Kimia (Feature Importance)", level=2)
    doc.add_paragraph(
        "Berdasarkan analisis Random Forest, tiga atribut paling berpengaruh dalam menentukan kualitas produk minuman anggur adalah:\n"
        "1. Alcohol (Persentase Kadar Alkohol): Skor 0.165\n"
        "2. Sulphates (Kadar Senyawa Sulfat): Skor 0.128\n"
        "3. Volatile Acidity (Keasaman Mudah Menguap): Skor 0.119"
    )
    doc.add_picture('feature_importance.png', width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    
    # ---------------------------------------------------------
    # SECTION 6: TAUTAN SUMBER DAYA & KONTRIBUSI TIM
    # ---------------------------------------------------------
    h6 = doc.add_heading("6. Tautan Dataset, Colab, Video & Kontribusi Tim", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_heading("a. Tautan Sumber Daya Publik:", level=2)
    p = doc.add_paragraph()
    p.add_run("1. Link Dataset: ").bold = True
    p.add_run("https://archive.ics.uci.edu/ml/machine-learning-databases/wine-quality/winequality-red.csv\n")
    p.add_run("2. Link Google Colab (.ipynb): ").bold = True
    p.add_run("[ISIKAN LINK GOOGLE COLAB KAMU DI SINI]\n")
    p.add_run("3. Link Video Penjelasan (YouTube/Drive): ").bold = True
    p.add_run("[ISIKAN LINK VIDEO PRESENTASI / GOOGLE DRIVE DI SINI]")
    
    doc.add_heading("b. Tabel Kontribusi Anggota Kelompok:", level=2)
    kontrib_tbl = doc.add_table(rows=3, cols=3)
    kontrib_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    k_headers = ["Nama Anggota", "NIM", "Deskripsi Kontribusi Pekerjaan"]
    for j, h in enumerate(k_headers):
        cell = kontrib_tbl.rows[0].cells[j]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "003366")
        
    kontrib_data = [
        ("[NAMA ANGGOTA 1]", "[NIM ANGGOTA 1]", "Manajemen dataset, pembuatan script eksperimen Python, tuning hyperparameter Random Forest & SVM, visualisasi diagram."),
        ("[NAMA ANGGOTA 2]", "[NIM ANGGOTA 2]", "Penyusunan laporan Latar Belakang, pembahasan cara kerja algoritma, dokumentasi video penjelasan, dan upload Google Colab.")
    ]
    for idx, (n, nim_val, des) in enumerate(kontrib_data):
        row = kontrib_tbl.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(n).bold = True
        row.cells[1].paragraphs[0].add_run(nim_val)
        row.cells[2].paragraphs[0].add_run(des)
        bg = "F9FAFB" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(row.cells[0], bg)
        set_cell_background(row.cells[1], bg)
        set_cell_background(row.cells[2], bg)
        
    # Save document
    doc_path = "Laporan_UAS_Machine_Learning.docx"
    doc.save(doc_path)
    print(f"[SUCCESS] Document successfully saved as {doc_path}")

if __name__ == '__main__':
    create_report()
